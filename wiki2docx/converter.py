"""HTML → DOCX converter for RoboMaster Wiki content.

Parses the Slate-editor HTML captured by the scraper and converts it into
a python-docx Document with full formatting support:

- Headings (h1–h6) mapped to Word heading styles
- Tables (<table>, <tr>, <td>, <th>)
- Lists (<ul>/<ol> with <li>)
- Images (<img>) — downloaded and embedded
- Inline formatting (bold, italic, underline, color, font-size, links)
- Paragraph alignment (text-align from style attribute)
- Page breaks between top-level sections
- Title page with configurable title/subtitle
- Table of contents from the section hierarchy
"""

import json
import os
import re
import ssl
import tempfile
import urllib.request

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# =============================================================================
# Public API
# =============================================================================

def build_docx(
    data: dict,
    output_path: str,
    title_override: str = None,
    subtitle_override: str = None,
    font_name: str = "微软雅黑",
) -> str:
    """Convert scraped wiki data to a Word document.

    Parameters
    ----------
    data : dict
        The JSON data from the scraper (or loaded from a JSON file).
    output_path : str
        Where to write the .docx file.
    title_override : str, optional
        Title text to use on the cover page.
    subtitle_override : str, optional
        Subtitle text to use on the cover page.
    font_name : str
        Primary font for body text and headings.

    Returns
    -------
    str
        The absolute path to the created file.
    """
    doc = _make_document(font_name)

    sections = data["sections"]

    # Title page
    title_text = title_override or data.get("title", "")
    subtitle_text = subtitle_override or ""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    if subtitle_text:
        p2 = doc.add_paragraph(f"\n{subtitle_text}")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p2.runs:
            r.font.size = Pt(14)

    doc.add_paragraph("\n")

    # Table of contents
    doc.add_heading("目  录", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    _build_toc(doc, sections)
    doc.add_page_break()

    # Content
    converter = _HtmlDocxConverter(doc, font_name)

    for idx, section in enumerate(sections):
        title = section["title"].strip()
        html = section.get("content_html", "")
        text = section.get("content", "")
        depth = section["level"]

        if not html and not text:
            continue

        heading_level = min(depth + 1, 3)
        doc.add_heading(title, level=heading_level)

        if html:
            converter.parse(html)
        elif text:
            for line in text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # Page break between top-level sections
        if idx + 1 < len(sections) and sections[idx + 1]["level"] == 0:
            doc.add_page_break()

    output_path = os.path.abspath(output_path)
    doc.save(output_path)
    return output_path


# =============================================================================
# Internal helpers
# =============================================================================

def _make_document(font_name: str) -> Document:
    """Create a Document with base styles configured."""
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.5)

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5
    _set_east_asian_font(normal, font_name)

    # Heading styles
    heading_sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13), 4: Pt(12)}
    for i in range(1, 5):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = RGBColor(0, 51, 102)
        hs.font.name = font_name
        hs.font.size = heading_sizes.get(i, Pt(12))
        _set_east_asian_font(hs, font_name)

    return doc


def _set_east_asian_font(style, font_name):
    """Ensure East-Asian glyphs use the correct font."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _build_toc(doc, sections):
    """Render a table-of-contents from the section list."""
    for s in sections:
        title = s["title"].strip()
        depth = s["level"]
        has_content = len(s.get("content", "")) > 0

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(depth * 1.2)

        run = p.add_run(title)
        run.font.size = Pt(12) if depth == 0 else Pt(10.5)
        run.font.bold = depth == 0

        if not has_content:
            run.font.color.rgb = RGBColor(160, 160, 160)


# =============================================================================
# HTML → DOCX converter
# =============================================================================

class _HtmlDocxConverter:
    """Parse the Slate-editor HTML and produce Word paragraphs / tables / lists."""

    def __init__(self, doc: Document, font_name: str):
        self.doc = doc
        self.font_name = font_name
        self.cur_p = None
        self.cur_heading = None
        self.cur_table_data = []  # list of list of cell dicts
        self.cur_row = []
        self.cell_text = ""
        self.in_cell = False
        self.cell_is_header = False
        self.in_list = False
        self.list_type = None
        self.list_counter = 0
        self.stack = []
        self.style = {
            "bold": False,
            "italic": False,
            "underline": False,
            "size": None,
            "color": None,
            "link": None,
        }

    # ---- Public ----

    def parse(self, html: str):
        """Parse and convert an HTML string into Word content."""
        html = self._clean_slate_noise(html)
        if not html.strip():
            return

        tag_pat = re.compile(
            r"<(/?)(\w+)((?:\s+[^>]*?)?)\s*(/?)>", re.IGNORECASE
        )
        pos = 0

        for m in tag_pat.finditer(html):
            self._emit_text(html[pos : m.start()])
            self._dispatch(
                is_close=bool(m.group(1)),
                tag=m.group(2).lower(),
                attrs=dict(re.findall(r"""(\S+)=["']([^"']*)["']""", m.group(3))),
                self_close=bool(m.group(4)),
            )
            pos = m.end()

        self._emit_text(html[pos:])

        # If we were building a table that never closed, finalize it
        if self.cur_table_data:
            self._finish_table()

    # ---- Tag dispatch ----

    def _dispatch(self, is_close, tag, attrs, self_close):
        if self.skip_content and not is_close:
            return

        # Self-closing
        if self_close:
            if tag == "br":
                self._add_br()
            elif tag == "img":
                self._add_image(attrs.get("src", ""))
            elif tag == "hr":
                self._add_hr()
            return

        # Closing
        if is_close:
            return self._close(tag)

        # Opening
        self.stack.append(tag)

        handler = getattr(self, f"_open_{tag}", None)
        if handler:
            handler(attrs)
        else:
            # Unknown tag — treat as inline container
            pass

    def _close(self, tag):
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.cur_heading = None
        elif tag in ("p", "div", "li"):
            self.cur_p = None
        elif tag in ("strong", "b"):
            self.style["bold"] = False
        elif tag in ("em", "i"):
            self.style["italic"] = False
        elif tag == "u":
            self.style["underline"] = False
        elif tag == "a":
            self.style["link"] = None
        elif tag == "span":
            self.style["size"] = None
            self.style["color"] = None
        elif tag in ("ul", "ol"):
            self.in_list = False
            self.list_type = None
            self.list_counter = 0
        elif tag == "table":
            self._finish_table()
        elif tag == "tr":
            if self.cur_row:
                self.cur_table_data.append(self.cur_row)
            self.cur_row = []
        elif tag in ("td", "th"):
            self.cur_row.append(
                {
                    "text": self.cell_text,
                    "is_header": self.cell_is_header,
                }
            )
            self.cell_text = ""
            self.in_cell = False
        elif tag in ("thead", "tbody"):
            pass

        # Pop stack
        if tag in self.stack:
            while self.stack and self.stack[-1] != tag:
                self.stack.pop()
            if self.stack:
                self.stack.pop()

    # ---- Opening handlers for structural tags ----

    def _open_h1(self, attrs):
        self._open_heading(1, attrs)

    def _open_h2(self, attrs):
        self._open_heading(2, attrs)

    def _open_h3(self, attrs):
        self._open_heading(3, attrs)

    def _open_h4(self, attrs):
        self._open_heading(4, attrs)

    def _open_h5(self, attrs):
        self._open_heading(5, attrs)

    def _open_h6(self, attrs):
        self._open_heading(6, attrs)

    def _open_heading(self, level, attrs):
        self.cur_p = None
        self.cur_heading = self.doc.add_heading(level=level)
        align = _extract_align(attrs.get("style", ""))
        if align is not None:
            self.cur_heading.alignment = align

    def _open_p(self, attrs):
        self.cur_heading = None
        self.cur_p = self.doc.add_paragraph()
        align = _extract_align(attrs.get("style", ""))
        if align is not None:
            self.cur_p.alignment = align

    def _open_div(self, attrs):
        self.cur_heading = None
        self.cur_p = self.doc.add_paragraph()
        align = _extract_align(attrs.get("style", ""))
        if align is not None:
            self.cur_p.alignment = align

    def _open_table(self, attrs):
        self.cur_table_data = []

    def _open_tr(self, attrs):
        self.cur_row = []

    def _open_td(self, attrs):
        self.in_cell = True
        self.cell_is_header = False
        self.cell_text = ""

    def _open_th(self, attrs):
        self.in_cell = True
        self.cell_is_header = True
        self.cell_text = ""

    def _open_ul(self, attrs):
        self.in_list = True
        self.list_type = "ul"
        self.list_counter = 0

    def _open_ol(self, attrs):
        self.in_list = True
        self.list_type = "ol"
        self.list_counter = 0

    def _open_li(self, attrs):
        self.cur_heading = None
        self.cur_p = self.doc.add_paragraph()
        self.cur_p.style = self.doc.styles["List Bullet"]

        if self.list_type == "ol":
            self.list_counter += 1
            prefix = f"{self.list_counter}. "
        else:
            prefix = "• "

        run = self.cur_p.add_run(prefix)
        run.font.size = Pt(10.5)

    # ---- Inline formatting ----

    def _open_strong(self, attrs):
        self.style["bold"] = True

    def _open_b(self, attrs):
        self.style["bold"] = True

    def _open_em(self, attrs):
        self.style["italic"] = True

    def _open_i(self, attrs):
        self.style["italic"] = True

    def _open_u(self, attrs):
        self.style["underline"] = True

    def _open_a(self, attrs):
        self.style["link"] = attrs.get("href", "")

    def _open_span(self, attrs):
        sty = attrs.get("style", "")
        if not sty:
            return
        cm = re.search(r"color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)", sty)
        if cm:
            self.style["color"] = tuple(int(g) for g in cm.groups())
        fm = re.search(r"font-size:\s*(\d+)px", sty)
        if fm:
            self.style["size"] = int(fm.group(1)) * 0.75

    # ---- Text emission ----

    def _emit_text(self, text):
        text = _unescape(text).replace("﻿", "")
        if not text:
            return

        if self.in_cell:
            self.cell_text += text
            return

        target = self.cur_heading or self.cur_p
        if target is not None:
            run = target.add_run(text)
            self._apply_style(run)
        else:
            # Orphan text — create a new paragraph
            self.cur_p = self.doc.add_paragraph()
            run = self.cur_p.add_run(text)
            self._apply_style(run)

    def _add_br(self):
        if self.in_cell:
            self.cell_text += "\n"
            return
        target = self.cur_heading or self.cur_p
        if target is not None:
            target.add_run("\n")

    def _add_hr(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "999999")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_image(self, src):
        if not src:
            return
        p = self.cur_p or self.doc.add_paragraph()
        try:
            ctx = ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = ssl.CERT_NONE
            req = urllib.request.Request(src, headers={"User-Agent": "Mozilla/5.0"})
            resp = urllib.request.urlopen(req, context=ctx, timeout=10)
            data = resp.read()
            if data:
                ext = os.path.splitext(src.split("?")[0])[1] or ".png"
                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tf:
                    tf.write(data)
                run = p.add_run()
                run.add_picture(tf.name, width=Inches(5))
                os.unlink(tf.name)
        except Exception:
            run = p.add_run(f"[图片: {src}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.size = Pt(9)

    def _apply_style(self, run):
        run.font.name = self.font_name
        if self.style["bold"]:
            run.font.bold = True
        if self.style["italic"]:
            run.font.italic = True
        if self.style["underline"]:
            run.font.underline = True
        if self.style["size"]:
            run.font.size = Pt(self.style["size"])
        if self.style["color"]:
            run.font.color.rgb = RGBColor(*self.style["color"])
        if self.style["link"]:
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.font.underline = True

    # ---- Table finalization ----

    def _finish_table(self):
        if not self.cur_table_data:
            return

        num_cols = max(len(r) for r in self.cur_table_data)
        num_rows = len(self.cur_table_data)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row_data in enumerate(self.cur_table_data):
            for ci, cell_data in enumerate(row_data):
                if ci >= num_cols:
                    break
                cell = table.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_data["text"].strip())
                run.font.name = self.font_name
                run.font.size = Pt(10)
                if cell_data.get("is_header"):
                    run.font.bold = True
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "003366")
                    shading.set(qn("w:val"), "clear")
                    cell._tc.get_or_add_tcPr().append(shading)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        self.doc.add_paragraph("")
        self.cur_table_data = []
        self.cur_row = []

    # ---- Helpers ----

    @staticmethod
    def _clean_slate_noise(html: str) -> str:
        patterns = [
            r'\bdata-slate-[a-zA-Z-]+(?:="[^"]*")?',
            r'\bdata-[a-z]+-[a-z-]+(?:="[^"]*")?',
            r'\bsuppresscontenteditablewarning=""',
            r'\bcontenteditable="false"',
            r'\brole="[^"]*"',
            r'\bspellcheck="[^"]*"',
            r'\bautocorrect="[^"]*"',
            r'\bautocapitalize="[^"]*"',
            r'\bdraggable="[^"]*"',
            r'\bid=""',
            r'\bclass="[^"]*"',
            r'\bstyle=""',
        ]
        for pat in patterns:
            html = re.sub(pat, "", html)
        return html

    @property
    def skip_content(self):
        return False


# =============================================================================
# Shared utilities
# =============================================================================

def _unescape(s: str) -> str:
    return (
        s.replace("&nbsp;", " ")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&amp;", "&")
        .replace("&quot;", '"')
        .replace("&#x27;", "'")
    )


def _extract_align(style_str: str):
    m = re.search(r"text-align:\s*(left|center|right|justify)", style_str)
    if not m:
        return None
    return {"left": 0, "center": 1, "right": 2, "justify": 3}.get(m.group(1))
